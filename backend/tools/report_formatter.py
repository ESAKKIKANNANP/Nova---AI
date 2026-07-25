# =============================================================================
# backend/tools/report_formatter.py
#
# Formats the structured JSON output into Markdown and PDF.
# =============================================================================

import os
import logging
import markdown
from fpdf import FPDF
from docx import Document
from typing import Dict, Any, List

log = logging.getLogger(__name__)

class ReportFormatter:
    def __init__(self, artifact_dir: str = "artifacts/reports"):
        self.artifact_dir = artifact_dir
        os.makedirs(self.artifact_dir, exist_ok=True)

    def generate_markdown(self, session_id: str, insight_data: Dict[str, Any], plot_paths: List[str] = None) -> str:
        """
        Generates a beautiful Markdown document from the structured JSON.
        """
        md = f"# Business Insight & Action Report (Session: {session_id})\n\n"
        
        md += "## Executive Summary\n"
        md += f"{insight_data.get('executive_summary', 'No summary provided.')}\n\n"
        
        md += "## Key Business Insights\n"
        for insight in insight_data.get("business_insights", []):
            md += f"- {insight}\n"
        md += "\n"
        
        md += "## Strategic Recommendations\n"
        for rec in insight_data.get("recommendations", []):
            md += f"- {rec}\n"
        md += "\n"
        
        md += "## Risk Analysis\n"
        for risk in insight_data.get("risk_analysis", []):
            md += f"- **Risk**: {risk}\n"
        md += "\n"
        
        md += "## Next Steps\n"
        for step in insight_data.get("next_steps", []):
            md += f"- {step}\n"
        md += "\n"
        
        if plot_paths:
            md += "## Appendix: Supporting Visualizations\n"
            for path in plot_paths:
                # Assuming plots are local for now
                file_name = os.path.basename(path)
                md += f"![{file_name}]({path})\n\n"
                
        # Save to disk
        md_path = os.path.join(self.artifact_dir, f"report_{session_id}.md")
        with open(md_path, "w") as f:
            f.write(md)
            
        return md_path

    def generate_pdf(self, md_path: str) -> str:
        """
        Generates a basic PDF from the markdown text using fpdf2.
        For MVP, we will parse out basic headers and bullets.
        """
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        pdf.set_font("Helvetica", size=12)
        
        if not os.path.exists(md_path):
            log.error(f"Markdown file not found: {md_path}")
            return None
            
        with open(md_path, "r") as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            if not line:
                pdf.ln(5)
                continue
                
            def clean_txt(s: str) -> str:
                return s.replace('\u2022', '*').replace('•', '*').encode('latin-1', 'replace').decode('latin-1')

            if line.startswith("# "):
                pdf.set_font("Helvetica", style="B", size=18)
                pdf.multi_cell(0, 10, text=clean_txt(line.replace("# ", "")))
                pdf.ln(10)
                pdf.set_font("Helvetica", size=12)
            elif line.startswith("## "):
                pdf.set_font("Helvetica", style="B", size=14)
                pdf.multi_cell(0, 10, text=clean_txt(line.replace("## ", "")))
                pdf.ln(10)
                pdf.set_font("Helvetica", size=12)
            elif line.startswith("- "):
                # Basic bold parsing for Risk Analysis
                txt = line.replace("- ", "* ")
                if "**Risk**:" in txt:
                    txt = txt.replace("**Risk**:", "Risk:")
                pdf.multi_cell(0, 8, text=clean_txt(txt))
                pdf.ln(8)
            elif line.startswith("!["):
                # Skip images in the basic PDF generator for MVP to avoid fpdf image pathing issues
                pass
            else:
                pdf.multi_cell(0, 6, text=clean_txt(line))
                pdf.ln(6)
                
        pdf_path = md_path.replace(".md", ".pdf")
        pdf.output(pdf_path)
        
        return pdf_path

    def generate_html(self, md_path: str) -> str:
        """
        Converts Markdown to styled HTML.
        """
        if not os.path.exists(md_path):
            log.error(f"Markdown file not found: {md_path}")
            return None
            
        with open(md_path, "r") as f:
            md_content = f.read()
            
        html_content = markdown.markdown(md_content)
        
        # Add basic styling
        html_styled = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; color: #333; }}
                h1 {{ color: #2c3e50; }}
                h2 {{ color: #34495e; border-bottom: 1px solid #eee; padding-bottom: 5px; }}
                img {{ max-width: 100%; height: auto; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        html_path = md_path.replace(".md", ".html")
        with open(html_path, "w") as f:
            f.write(html_styled)
            
        return html_path

    def generate_docx(self, md_path: str) -> str:
        """
        Converts basic Markdown headers and lists to a DOCX file using python-docx.
        """
        if not os.path.exists(md_path):
            log.error(f"Markdown file not found: {md_path}")
            return None
            
        doc = Document()
        
        with open(md_path, "r") as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith("# "):
                doc.add_heading(line.replace("# ", ""), level=1)
            elif line.startswith("## "):
                doc.add_heading(line.replace("## ", ""), level=2)
            elif line.startswith("- "):
                txt = line.replace("- ", "")
                if "**Risk**:" in txt:
                    txt = txt.replace("**Risk**:", "Risk:")
                doc.add_paragraph(txt, style='List Bullet')
            elif line.startswith("!["):
                # Skip images for DOCX MVP to avoid complex embedding logic
                pass
            else:
                doc.add_paragraph(line)
                
        docx_path = md_path.replace(".md", ".docx")
        doc.save(docx_path)
        
        return docx_path
